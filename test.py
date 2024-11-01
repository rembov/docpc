import os
from docx import Document
import fitz  # PyMuPDF для работы с PDF
from openpyxl import load_workbook

# Путь к директории и справочному файлу
directory = "C:/Users/user/Desktop/123"
reference_file = "C:/Users/user/Desktop/123/опись.docx"  # Замените на ваш файл


def load_reference_data(reference_path):
    """
    Загружает справочные данные из файла .docx и возвращает их в виде списка словарей.
    :param reference_path: Путь к файлу справочника
    :return: Список словарей с данными справочника
    """
    try:
        if not os.path.exists(reference_path):
            print(f"Ошибка: файл справочника '{reference_path}' не найден.")
            return []

        document = Document(reference_path)
        keys = ["Наименование:", "Обозначение:", "Количество листов:", "Формат:"]
        reference_data = []

        entry = {}
        for para in document.paragraphs:
            text = para.text.strip()
            for key in keys:
                if text.startswith(key):
                    entry[key] = text[len(key):].strip()
                    # Если все ключи найдены, добавляем запись в справочные данные
                    if len(entry) == len(keys):
                        reference_data.append(entry)
                        entry = {}  # Сброс для следующего файла

        return reference_data
    except Exception as e:
        print(f"Ошибка загрузки справочника {reference_path}: {str(e)}")
        return []


def get_file_metadata(file_path):
    """
    Извлекает метаданные файла (наименование, обозначение, количество листов, формат).
    :param file_path: Путь к файлу
    :return: Словарь с метаданными файла
    """
    ext = file_path.split('.')[-1].lower()
    file_metadata = {
        "Наименование:": os.path.basename(file_path),
        "Обозначение:": "Не определено",
        "Количество листов:": 0,
        "Формат:": ext
    }

    try:
        if ext == "pdf":
            with fitz.open(file_path) as pdf:
                file_metadata["Количество листов:"] = pdf.page_count
        elif ext == "docx":
            doc = Document(file_path)
            file_metadata["Обозначение:"] = doc.core_properties.subject or "Не определено"
            file_metadata["Количество листов:"] = len(doc.paragraphs)
        elif ext == "txt":
            with open(file_path, "r", encoding="utf-8") as file:
                content = file.read()
                file_metadata["Количество листов:"] = content.count('\n') // 50 + 1
        elif ext == "xlsx":
            workbook = load_workbook(file_path, data_only=True)
            file_metadata["Количество листов:"] = len(workbook.sheetnames)
    except Exception as e:
        print(f"Ошибка извлечения метаданных из файла {file_path}: {str(e)}")
        return None

    return file_metadata


def compare_with_directory(reference_data, directory):
    """
    Сравнивает данные справочника с файлами в каталоге.
    :param reference_data: Справочные данные
    :param directory: Путь к каталогу с документами
    """
    for ref_entry in reference_data:
        ref_name = ref_entry.get("Наименование:")
        if not ref_name:
            continue

        file_path = os.path.join(directory, ref_name)

        # Если файл существует, выполняем сравнение, иначе добавляем запись о его отсутствии
        if os.path.isfile(file_path):
            file_metadata = get_file_metadata(file_path)
            if not file_metadata:
                continue

            # Сравнение характеристик
            print(f"\nПроверка файла: {ref_name}")
            for key, ref_value in ref_entry.items():
                file_value = file_metadata.get(key, "Не определено")
                if ref_value != file_value:
                    print(f"{key.capitalize()} отличается: Справочник ({ref_value}) vs Файл ({file_value})")
                else:
                    print(f"{key.capitalize()} совпадает.")
        else:
            print(f"\nФайл '{ref_name}' отсутствует в каталоге.")


def main():
    # Загрузка справочных данных
    reference_data = load_reference_data(reference_file)
    if not reference_data:
        print("Не удалось загрузить данные справочника.")
        return

    # Сравнение справочника с файлами в каталоге
    compare_with_directory(reference_data, directory)


if __name__ == "__main__":
    main()
