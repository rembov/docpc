import zipfile
import re
from difflib import ndiff
import PyPDF2
import docx
import rarfile
import py7zr
import pytesseract
import fitz  # PyMuPDF
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
from openpyxl import load_workbook
from PIL import Image, ImageDraw
import os
import logging
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
import shutil

output_directory_text_images = ""
output_directory_numbering = ""
# Настройка логирования
logging.basicConfig(filename="process.txt", level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

def extract_data_from_pdf(pdf_path, output_dir):
    """
    Извлекает текст и изображения из PDF.
    :param pdf_path: Путь к PDF файлу
    :param output_dir: Директория для сохранения изображений
    :return: Извлечённый текст
    """
    try:
        data = ""
        with fitz.open(pdf_path) as pdf_file:
            for i, page in enumerate(pdf_file):
                data += page.get_text()
                # Извлечение изображений
                images = page.get_images(full=True)
                for img_index, img in enumerate(images):
                    xref = img[0]
                    pix = fitz.Pixmap(pdf_file, xref)
                    if pix.n > 4:  # если изображение в формате CMYK, конвертируем в RGB
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    image_path = os.path.join(output_dir, f"page_{i + 1}_image_{img_index + 1}.png")
                    pix.save(image_path)
                    logging.info(f"Изображение сохранено: {image_path}")
                    pix = None  # освобождение памяти
        logging.info(f"Текст успешно извлечён из {pdf_path}")
        return data
    except Exception as e:
        logging.error(f"Ошибка при извлечении текста из {pdf_path}: {str(e)}")
        return ""


def extract_data_from_docx(docx_path, output_dir):
    """
    Извлекает текст и изображения из документа Word (.docx).
    :param docx_path: Путь к документу Word
    :param output_dir: Директория для сохранения изображений
    :return: Извлечённый текст
    """
    try:
        doc = Document(docx_path)
        data = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        # Извлечение изображений
        for rel in doc.rels.values():
            if "image" in rel.target_ref:
                img = rel.target_part.blob
                image_path = os.path.join(output_dir, f"{rel.target_ref.split('/')[-1]}")
                with open(image_path, "wb") as f:
                    f.write(img)
                logging.info(f"Изображение сохранено: {image_path}")

        logging.info(f"Текст успешно извлечён из {docx_path}")
        return data
    except Exception as e:
        logging.error(f"Ошибка при извлечении текста из {docx_path}: {str(e)}")
        return ""


def extract_data_from_txt(txt_path):
    """
    Извлекает текст из текстового файла (.txt).
    :param txt_path: Путь к текстовому файлу
    :return: Извлечённый текст
    """
    try:
        with open(txt_path, "r", encoding="utf-8") as file:
            data = file.read()
        logging.info(f"Текст успешно извлечён из {txt_path}")
        return data
    except Exception as e:
        logging.error(f"Ошибка при извлечении текста из {txt_path}: {str(e)}")
        return ""


def select_output_directory_for_text_images():
    global output_directory_text_images
    output_directory_text_images = filedialog.askdirectory(title="Выберите директорию для вывода текста и изображений")
    if output_directory_text_images:
        logging.info(f"Директория для текста и изображений установлена: {output_directory_text_images}")


def select_output_directory_for_numbering():
    global output_directory_numbering
    output_directory_numbering = filedialog.askdirectory(title="Выберите директорию для вывода с нумерацией")
    if output_directory_numbering:
        logging.info(f"Директория для нумерации установлена: {output_directory_numbering}")


def extract_data_from_xlsx(xlsx_path):
    """
    Извлекает текст из Excel (.xlsx).
    :param xlsx_path: Путь к Excel файлу
    :return: Извлечённый текст
    """
    try:
        workbook = load_workbook(xlsx_path, data_only=True)
        data = ""
        for sheet in workbook.sheetnames:
            sheet_data = "\n".join(
                ["\t".join([str(cell.value) if cell.value else "" for cell in row]) for row in workbook[sheet].rows]
            )
            data += f"\nЛист {sheet}:\n{sheet_data}\n"
        logging.info(f"Текст успешно извлечён из {xlsx_path}")
        return data
    except Exception as e:
        logging.error(f"Ошибка при извлечении текста из {xlsx_path}: {str(e)}")
        return ""


def extract_archive(file_path, extract_to):
    """
    Извлекает файлы из архива в указанную директорию.

    :param file_path: Путь к архивному файлу
    :param extract_to: Директория для извлечения
    :return: True, если извлечение прошло успешно, иначе False
    """
    if not file_path or not extract_to:
        logging.error("Необходимо указать расположение архива и место для разархивации.")
        messagebox.showerror("Ошибка", "Необходимо указать расположение архива и место для разархивации.")
        return False

    try:
        # Создаём директорию для извлечения, если она не существует
        os.makedirs(extract_to, exist_ok=True)

        if file_path.endswith('.zip'):
            with zipfile.ZipFile(file_path, 'r') as archive:
                archive.extractall(extract_to)  # Извлечение всех файлов из ZIP
                logging.info(f"Архив {file_path} успешно извлечён в {extract_to}.")

        elif file_path.endswith('.rar'):
            # Проверка наличия unrar
            rarfile.UNRAR_TOOL = "path/to/unrar"  # Укажите путь к unrar, если он не в PATH
            with rarfile.RarFile(file_path, 'r') as archive:
                archive.extractall(extract_to)  # Извлечение всех файлов из RAR
                logging.info(f"Архив {file_path} успешно извлечён в {extract_to}.")

        elif file_path.endswith('.7z'):
            with py7zr.SevenZipFile(file_path, mode='r') as archive:
                archive.extractall(extract_to)  # Извлечение всех файлов из 7z
                logging.info(f"Архив {file_path} успешно извлечён в {extract_to}.")

        else:
            logging.error("Неподдерживаемый формат архива.")
            messagebox.showerror("Ошибка", "Неподдерживаемый формат архива.")
            return False

        messagebox.showinfo("Успех", f"Архив успешно извлечён в {extract_to}")
        return True

    except rarfile.Error as e:
        logging.error(f"Ошибка при извлечении RAR-архива {file_path}: {str(e)}")
        messagebox.showerror("Ошибка", f"Не удалось извлечь RAR-архив. {str(e)}")
        return False

    except (zipfile.BadZipFile, py7zr.Bad7zFile) as e:
        logging.error(f"Ошибка: архив повреждён или имеет неверный формат {file_path}: {str(e)}")
        messagebox.showerror("Ошибка", "Архив повреждён или имеет неверный формат.")
        return False

    except PermissionError:
        logging.error(f"Ошибка: недостаточно прав для записи в {extract_to}.")
        messagebox.showerror("Ошибка", "Недостаточно прав для записи в указанную директорию.")
        return False

    except Exception as e:
        logging.error(f"Ошибка извлечения файла {file_path}: {str(e)}")
        messagebox.showerror("Ошибка", "Ошибка при извлечении архива.")
        return False

def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with fitz.open(pdf_path) as pdf:
            for page_num in range(len(pdf)):
                page = pdf[page_num]
                text += page.get_text("text")
                if not text.strip():
                    pix = page.get_pixmap()
                    image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    text += pytesseract.image_to_string(image)
    except Exception as e:
        logging.error(f"Ошибка извлечения текста из PDF {pdf_path}: {str(e)}")
    return text


def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        logging.error(f"Ошибка извлечения текста из DOCX {docx_path}: {str(e)}")
    return ""


def extract_data_from_excel(xlsx_path):
    try:
        workbook = load_workbook(xlsx_path, data_only=True)
        sheet = workbook.active
        data = [[cell.value for cell in row] for row in sheet.iter_rows()]
        return data
    except Exception as e:
        logging.error(f"Ошибка извлечения данных из Excel {xlsx_path}: {str(e)}")
    return []


def compare_with_reference(data, reference_path):
    try:
        reference = pd.read_excel(reference_path)
        matched_data = [item for item in data if item in reference.values]
        return matched_data
    except Exception as e:
        logging.error(f"Ошибка при загрузке справочника {reference_path}: {str(e)}")
    return []


def extract_file_metadata(file_path):
    """
    Извлекает метаданные о документе.
    :param file_path: Путь к файлу
    :return: Кортеж с наименованием, обозначением, количеством страниц и форматом
    """
    # Получаем имя файла и его формат
    name = os.path.basename(file_path)
    format = os.path.splitext(file_path)[1][1:].lower()  # Формат файла без точки

    designation = "Обозначение документа"  # Заглушка, заменить на реальное значение, если доступно
    pages = 0  # Количество страниц, инициализируем как 0

    try:
        # Обработка файлов PDF
        if format == 'pdf':
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                pages = len(reader.pages)  # Получаем количество страниц в PDF

        # Обработка файлов DOCX
        elif format == 'docx':
            doc = docx.Document(file_path)
            pages = len(doc.element.xpath('//w:sectPr'))  # Пример получения количества страниц

        # Обработка текстовых файлов
        elif format == 'txt':
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
                pages = content.count('\n') // 50 + 1  # Примерное количество страниц

        # Обработка изображений (например, сканированных документов)
        elif format in ['jpg', 'jpeg', 'png']:
            img = Image.open(file_path)
            text = pytesseract.image_to_string(img)  # Извлечение текста из изображения
            pages = 1  # Для изображений, можно считать 1 страницу, если изображение одно

    except Exception as e:
        logging.error(f"Ошибка при извлечении метаданных из файла {file_path}: {str(e)}")

    return name, designation, pages, format


def create_inventory(matched_data, output_path):
    """
    Создает опись документов в формате .docx и сохраняет её.
    :param matched_data: Данные для внесения в опись
    :param output_path: Путь для сохранения описи
    """
    # Создание документа Word с описью
    try:
        doc = Document()

        for document in matched_data:
            doc.add_paragraph(
                f"Наименование: {document['name']}\n"
                f"Обозначение: {document['designation']}\n"
                f"Количество листов: {document['pages']}\n"
                f"Формат: {document['format']}"
            )

        doc.save(output_path)
        logging.info(f"Опись успешно сохранена в {output_path}")
        messagebox.showinfo("Успех", "Опись успешно создана.")

    except Exception as e:
        logging.error(f"Ошибка при создании описи: {str(e)}")
        messagebox.showerror("Ошибка", "Ошибка при создании описи.")


def extract_data_from_documents(directory):
    """
    Извлекает данные о всех файлах из указанной директории и возвращает список документов.
    :param directory: Путь к директории с документами
    :return: Список данных о документах
    """
    documents = []

    # Проходим по всем файлам в директории
    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        if os.path.isfile(file_path):
            # Извлекаем метаданные файла
            name, designation, pages, format = extract_file_metadata(file_path)
            documents.append({
                'name': name,
                'designation': designation,
                'pages': pages,
                'format': format
            })

    return documents




def apply_number_to_file(file_path, number, output_path):
    ext = os.path.splitext(file_path)[-1].lower()

    try:
        # Нанесение номера на .docx
        if ext == ".docx":
            doc = Document(file_path)
            doc.add_paragraph(f"Номер: {number}")
            doc.save(output_path)

        # Нанесение номера на .pdf
        elif ext == ".pdf":
            # Используем временный файл для предотвращения ошибки "save to original must be incremental"
            temp_output_path = output_path + "_temp.pdf"
            doc = fitz.open(file_path)
            first_page = doc[0]
            first_page.insert_text((10, 10), f"{number}", fontsize=12)
            doc.save(temp_output_path)
            doc.close()
            # Перемещаем временный файл на место оригинального output_path
            shutil.move(temp_output_path, output_path)

        # Нанесение номера на .txt
        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(f"Номер: {number}\n{content}")

        # Нанесение номера на .xlsx
        elif ext == ".xlsx":
            workbook = load_workbook(file_path)
            sheet = workbook.active
            sheet["A1"] = f"Номер: {number}"
            workbook.save(output_path)

        logging.info(f"Номер {number} успешно нанесен на файл {file_path} и сохранен как {output_path}")

    except Exception as e:
        logging.error(f"Ошибка нанесения номера на файл {file_path}: {str(e)}")
        messagebox.showerror("Ошибка", f"Ошибка при нанесении номера на файл {file_path}: {str(e)}")


def rename_file_with_dialog():
    try:
        current_file_path = filedialog.askopenfilename(title="Выберите файл для переименования")
        if not current_file_path:
            return

        new_name = simpledialog.askstring("Новое имя", "Введите новое имя для файла:")
        if not new_name:
            return

        save_directory = filedialog.askdirectory(title="Выберите директорию для сохранения")
        if not save_directory:
            return

        new_file_path = os.path.join(save_directory, new_name + os.path.splitext(current_file_path)[1])
        os.rename(current_file_path, new_file_path)
        logging.info(f"Файл {current_file_path} переименован и сохранен как {new_file_path}")
        messagebox.showinfo("Успех", f"Файл успешно переименован и перемещен в {new_file_path}")

    except Exception as e:
        logging.error(f"Ошибка при переименовании файла: {str(e)}")
        messagebox.showerror("Ошибка", "Ошибка при переименовании файла.")

def read_file_content(file_path, ext):
    """
    Считывает содержимое файла в зависимости от его расширения.
    :param file_path: Путь к файлу
    :param ext: Расширение файла
    :return: Содержимое файла в виде строки
    """
    content = ""
    try:
        if ext.lower() == ".txt":
            with open(file_path, 'r', encoding="utf-8", errors="ignore") as f:
                content = f.read()
        elif ext.lower() == ".docx":
            doc = Document(file_path)
            content = "\n".join([para.text for para in doc.paragraphs])
        elif ext.lower() == ".pdf":
            pdf_reader = PdfReader(file_path)
            content = "\n".join([page.extract_text() for page in pdf_reader.pages])
        logging.debug(f"Содержимое файла '{file_path}': {content[:100]}...")
    except Exception as e:
        logging.error(f"Ошибка при чтении содержимого файла '{file_path}': {str(e)}")

    return content
def check_and_rename_files(directory):
    """Проверяет и переименовывает файлы в каталоге, чтобы соответствовать определённому шаблону именования."""
    pattern = re.compile(r"^[A-Za-z0-9_-]+$")
    for filename in os.listdir(directory):
        if not pattern.match(filename):
            new_filename = re.sub(r'\W+', '_', filename)
            os.rename(os.path.join(directory, filename), os.path.join(directory, new_filename))
            logging.info(f"Файл переименован: {filename} -> {new_filename}")

class DocumentProcessorApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Document Processor")

        # Поля для путей
        self.archive_paths = tk.StringVar()
        self.reference_path = tk.StringVar()
        self.output_directory = tk.StringVar()
        self.numbers_path = tk.StringVar()
        self.files_directory = tk.StringVar()

        # Элементы интерфейса
        tk.Label(root, text="Путь к архивам:").grid(row=0, column=0, sticky="w")
        tk.Entry(root, textvariable=self.archive_paths, width=50).grid(row=0, column=1)
        tk.Button(root, text="Обзор", command=self.select_archives).grid(row=0, column=2)

        tk.Label(root, text="Путь к справочнику:").grid(row=3, column=0, sticky="w")
        tk.Entry(root, textvariable=self.reference_path, width=50).grid(row=3, column=1)
        tk.Button(root, text="Обзор", command=self.select_referenc1).grid(row=3, column=2)

        tk.Label(root, text="Директория для результатов:").grid(row=1, column=0, sticky="w")
        tk.Entry(root, textvariable=self.output_directory, width=50).grid(row=1, column=1)
        tk.Button(root, text="Обзор", command=self.select_output_directory).grid(row=1, column=2)

        tk.Label(root, text="Директория с файлами:").grid(row=2, column=0, sticky="w")
        tk.Entry(root, textvariable=self.files_directory, width=50).grid(row=2, column=1)
        tk.Button(root, text="Обзор", command=self.select_files_directory).grid(row=2, column=2)

        # Кнопки для функций с 3 строки и столбца
        button_commands = [
            (self.run_extraction, "Извлечь архив/архивы"),
            (self.run_extract_text_and_images, "Извлечь текст и изображения"),
            (self.run_inventory, "Сформировать опись"),
            (self.run_apply_numbers, "Нанести номера"),
            (self.run_rename_files, "Переименовать файлы"),
            (self.run_inventory_with_reference, "Опись со справочником\n+извлечение")
        ]

        for index, (command, text) in enumerate(button_commands):
            row = 5 + index // 3
            column = index % 3
            tk.Button(root, text=text, command=command).grid(row=row, column=column, padx=5, pady=5)

    def select_referenc1(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if file_path:
            self.reference_path.set(file_path)
            logging.info(f"Справочник выбран: {file_path}")

    def load_reference_from_excel(self,file_path):
        try:
            df = pd.read_excel(file_path)
            reference_dict = dict(zip(df.iloc[:, 0].astype(str), df.iloc[:, 1].astype(str)))  # Преобразуйте в строки
            print(reference_dict)
            return reference_dict
        except Exception as e:
            print(f"Ошибка при загрузке справочника: {e}")
            return {}
    def standardize_document_titles(self, documents, reference_dict):
        """
        Обновляет наименования документов на основе справочника.
        :param documents: Список словарей с метаданными документов
        :param reference_dict: Справочник с наименованиями
        :return: Обновленный список документов с эталонными наименованиями
        """
        for document in documents:
            original_name = document['name']
            for partial_name, full_name in reference_dict.items():
                if re.search(r'\b' + re.escape(partial_name) + r'\b', original_name, re.IGNORECASE):
                    document['name'] = full_name
                    logging.info(f"Наименование документа обновлено с '{original_name}' на '{full_name}'")
                    break
        return documents

    def extract_data_from_documents(self, directory, designation_dict):
        """
        Извлекает данные о документах из указанной директории.
        :param directory: Путь к директории с документами
        :return: Список данных о документах
        """
        documents = []
        for root, _, files in os.walk(directory):
            for filename in files:
                file_path = os.path.join(root, filename)
                ext = os.path.splitext(filename)[1].lower()
                if ext in ['.pdf', '.docx', '.txt']:
                    name, pages = self.extract_metadata(file_path, ext)
                    name1 = os.path.splitext(name)[0]
                    documents.append({
                        'name': name,
                        'designation': name1,
                        'pages': pages,
                        'format': ext[1:]
                    })
        return documents


    def extract_metadata(self, file_path, ext):
        """
        Извлекает метаданные из файла.
        :param file_path: Путь к файлу
        :param ext: Расширение файла
        :return: Наименование документа и количество страниц
        """
        name = os.path.basename(file_path)  # Имя файла
        pages = 0  # Количество страниц, инициализируем как 0

        try:
            if ext == '.pdf':
                from PyPDF2 import PdfReader
                with open(file_path, 'rb') as file:
                    reader = PdfReader(file)
                    pages = len(reader.pages)  # Получаем количество страниц в PDF

            elif ext == '.docx':
                doc = Document(file_path)
                pages = len(doc.element.xpath('//w:sectPr'))  # Пример получения количества страниц

            elif ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read()
                    pages = content.count('\n') // 50 + 1  # Примерное количество страниц
            elif ext == '.xlsx':
                workbook = load_workbook(filename=file_path, read_only=True)
                pages = len(workbook.sheetnames)

        except Exception as e:
            logging.error(f"Ошибка при извлечении метаданных из файла {file_path}: {str(e)}")

        return name, pages

    def create_inventory(self, documents, output_path):
        """
        Создает опись документов и сохраняет в формате .docx.
        :param documents: Данные о документах
        :param output_path: Путь для сохранения
        """
        try:
            doc = Document()
            for document in documents:
                doc.add_paragraph(
                    f"Наименование: {document['name']}\n"
                    f"Обозначение: {document['designation']}\n"
                    f"Количество листов: {document['pages']}\n"
                    f"Формат: {document['format']}"
                )
            doc.save(output_path)
            logging.info(f"Опись сохранена: {output_path}")
        except Exception as e:
            logging.error(f"Ошибка при создании описи: {e}")

    def run_inventory_with_reference(self):
        reference_path = self.reference_path.get()
        files_directory = self.files_directory.get()
        output_directory = self.output_directory.get()
        designation_file_path = self.reference_path.get()
        archive_paths = self.archive_paths.get()


        if not reference_path or not files_directory or not output_directory or not designation_file_path:
            messagebox.showerror("Ошибка", "Необходимо указать все пути.")
            return

        # Извлечение архивов
        self.run_extraction()
        # Нанесение номеров на файлы
        self.run_apply_numbers()

        # Загрузка справочника и обозначений
        reference_dict = self.load_reference_from_excel(reference_path)
        designation_dict = self.load_reference_from_excel(reference_path)

        # Извлечение данных о документах
        documents = self.extract_data_from_documents(files_directory,designation_dict)

        # Приведение наименований документов
        standardized_documents = self.standardize_document_titles(documents, reference_dict)
        #self.rename_files_according_to_reference(documents, reference_dict)
        # Нанесение номеров на документы
        self.rename_files_recursively(files_directory,reference_path)
        self.add_numbers_to_document_titles(standardized_documents)

        # Создание и сохранение описи
        output_path = os.path.join(output_directory, "опись.docx")
        self.create_inventory(standardized_documents, output_path)
        messagebox.showinfo("Успех", "Опись успешно создана с учетом справочника, обозначений и нанесенных номеров.")

    def rename_files_recursively(self, directory, reference_path):
        """
        Рекурсивно переименовывает файлы в каталоге, сравнивая их содержимое и название со справочником.
        :param directory: Путь к директории с файлами
        :param reference_path: Путь к Excel файлу справочника
        """
        try:
            # Загружаем справочник
            df_reference = pd.read_excel(reference_path)
            logging.debug(f"Загружены данные справочника:\n{df_reference.head()}")

            # Проверка структуры справочника
            if 'Русское название' not in df_reference.columns or 'Английское название' not in df_reference.columns:
                logging.error("Справочник должен содержать столбцы 'Русское название' и 'Английское название'.")
                return

            # Создание словаря для переименования (включаем и русские, и английские названия)
            reference_dict = {
                str(row['Русское название']).strip().lower(): str(row['Русское название']).strip()
                for _, row in df_reference.iterrows()
            }
            reference_dict.update({
                str(row['Английское название']).strip().lower(): str(row['Русское название']).strip()
                for _, row in df_reference.iterrows()
            })
            logging.debug(f"Словарь для переименования: {reference_dict}")

            # Проход по всем файлам в директории рекурсивно
            for root, _, files in os.walk(directory):
                logging.debug(f"Проверяем директорию: {root}")

                for filename in files:
                    file_path = os.path.join(root, filename)
                    name_without_ext, ext = os.path.splitext(filename)
                    name_lower = name_without_ext.lower()

                    logging.debug(f"Обрабатываем файл: {filename}")

                    # Проверка прав доступа
                    if not os.access(file_path, os.R_OK | os.W_OK):
                        logging.warning(f"Нет доступа к файлу: {file_path}")
                        continue

                    # Проверка совпадения по названию файла
                    new_name = reference_dict.get(name_lower)
                    if new_name:
                        new_filename = f"{new_name}{ext}"
                        new_file_path = os.path.join(root, new_filename)

                        # Переименовываем файл, если имя изменилось и файла с таким именем ещё нет
                        if new_file_path != file_path and not os.path.exists(new_file_path):
                            try:
                                os.rename(file_path, new_file_path)
                                logging.info(f"Файл '{filename}' переименован в '{new_filename}' по названию")
                            except Exception as e:
                                logging.error(f"Ошибка при переименовании файла '{filename}': {str(e)}")
                            continue

                    # Если совпадение по названию не найдено, проверяем содержимое файла
                    file_content = read_file_content(file_path, ext)
                    if not file_content:
                        logging.debug(f"Не удалось прочитать содержимое файла: {filename}")
                        continue

                    # Проверка совпадения по содержимому файла
                    for ref_name, new_name in reference_dict.items():
                        if re.search(rf"\b{re.escape(ref_name)}\b", file_content, re.IGNORECASE):
                            new_filename = f"{new_name}{ext}"
                            new_file_path = os.path.join(root, new_filename)

                            # Переименовываем файл по содержимому
                            if new_file_path != file_path and not os.path.exists(new_file_path):
                                try:
                                    os.rename(file_path, new_file_path)
                                    logging.info(f"Файл '{filename}' переименован в '{new_filename}' по содержимому")
                                    break
                                except Exception as e:
                                    logging.error(f"Ошибка при переименовании файла '{filename}': {str(e)}")
                            else:
                                logging.warning(f"Файл с именем '{new_filename}' уже существует.")
            logging.info("Рекурсивное переименование файлов завершено.")
        except Exception as e:
            logging.error(f"Ошибка при загрузке справочника или переименовании файлов: {str(e)}")

    def add_numbers_to_document_titles(self, documents):
        """
        Добавляет номера к наименованиям документов.
        :param documents: Список документов с метаданными
        :return: Обновленный список документов с добавленными номерами
        """
        for index, document in enumerate(documents, start=1):
            document['name'] = f"{index}. {document['name']}"  # Добавляем номер перед названием
            logging.info(f"Номер добавлен к документу: {document['name']}")
        return documents

    def select_archives(self):
            file_paths = filedialog.askopenfilenames(filetypes=[("Archive files", "*.zip *.rar *.7z")])
            if file_paths:
                self.archive_paths.set(";".join(file_paths))  # Store multiple paths as a semicolon-separated string
                logging.info(f"Архивы выбраны: {file_paths}")

    def run_extraction(self):
        """
        Метод для запуска процесса извлечения архивов.
        Извлекает архивы из указанных путей в заданную директорию.
        """
        archive_paths = self.archive_paths.get().split(";")  # Get multiple archive paths
        output_directory = self.output_directory.get()

        if not archive_paths or not output_directory:
            messagebox.showerror("Ошибка", "Необходимо указать архивы и директорию для извлечения.")
            return

        for archive_path in archive_paths:
            if not os.path.isfile(archive_path):
                messagebox.showerror("Ошибка", f"Указанный архив не существует: {archive_path}")
                continue

            if not os.path.exists(output_directory):
                messagebox.showerror("Ошибка", "Указанная директория для извлечения не существует.")
                return

            # Запускаем процесс извлечения
            if extract_archive(archive_path, output_directory):
                logging.info(f"Процесс извлечения для {archive_path} завершён успешно.")
            else:
                logging.error(f"Процесс извлечения для {archive_path} завершился с ошибкой.")

    def run_rename_files(self):
        """Запускает проверку и переименование файлов в указанной директории."""
        directory = self.files_directory.get()
        if directory:
            check_and_rename_files(directory)
            messagebox.showinfo("Успех", "Проверка и переименование файлов завершено.")
        else:
            messagebox.showerror("Ошибка", "Необходимо выбрать директорию с файлами для переименования.")


    def select_files_directory(self):
        """Метод для выбора директории с файлами для извлечения текста и изображений."""
        directory = filedialog.askdirectory()
        if directory:
            self.files_directory.set(directory)
            logging.info(f"Директория с файлами выбрана: {directory}")

    def select_reference(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.docx")])
        if file_path:
            self.reference_path.set(file_path)
            logging.info(f"Справочник выбран: {file_path}")

    def select_output_directory(self):
        directory = filedialog.askdirectory()
        if directory:
            self.output_directory.set(directory)
            logging.info(f"Директория для результатов выбрана: {directory}")

    def select_numbers(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if file_path:
            self.numbers_path.set(file_path)
            logging.info(f"Файл с номерами выбран: {file_path}")

    def run_extract_text_and_images(self):
        """Метод для извлечения текста и изображений из файлов различных форматов."""
        directory = self.files_directory.get()  # Использование новой директории
        if not directory:
            logging.error("Директория для извлечения не указана.")
            return

        # Запрос пути для сохранения извлечённого текста и изображений
        output_directory = filedialog.askdirectory(title="Выберите директорию для сохранения извлечённых данных")
        if not output_directory:
            logging.error("Директория для сохранения не указана.")
            return

        extracted_data = ""
        output_image_dir = os.path.join(output_directory, "extracted_images")  # Директория для сохранения изображений
        os.makedirs(output_image_dir, exist_ok=True)  # Создаём директорию, если она не существует

        # Путь для сохранения извлечённого текста
        output_text_path = os.path.join(output_directory, "extracted_data.txt")

        # Используем os.walk для рекурсивного обхода всех подкаталогов
        for root, _, files in os.walk(directory):
            for file_name in files:
                file_path = os.path.join(root, file_name)
                _, ext = os.path.splitext(file_name)

                if ext.lower() == ".xlsx":
                    extracted_data += f"\n--- Данные из {file_name} ---\n"
                    extracted_data += extract_data_from_xlsx(file_path)  # Прямо добавляем данные

                elif ext.lower() == ".docx":
                    extracted_data += f"\n--- Данные из {file_name} ---\n"
                    extracted_data += extract_data_from_docx(file_path,
                                                             output_image_dir)  # Передаём путь для сохранения изображений

                elif ext.lower() == ".txt":
                    extracted_data += f"\n--- Данные из {file_name} ---\n"
                    extracted_data += extract_data_from_txt(file_path)  # Измените, если нужно передать путь

                elif ext.lower() == ".pdf":
                    extracted_data += f"\n--- Данные из {file_name} ---\n"
                    extracted_data += extract_data_from_pdf(file_path,
                                                            output_image_dir)  # Передаём путь для сохранения изображений

        # Сохранение извлечённых данных в текстовый файл
        with open(output_text_path, "w", encoding="utf-8") as output_file:
            output_file.write(extracted_data)

        logging.info("Извлечение текста и изображений завершено.")
        messagebox.showinfo("Успех", "Извлечение завершено.")

    def get_all_files_in_directory(self,directory, extensions=('.pdf', '.docx', '.txt', '.xlsx')):
        """
        Рекурсивно находит все файлы с заданными расширениями в указанной директории и подкаталогах.
        :param directory: Путь к корневой директории для поиска
        :param extensions: Кортеж с расширениями файлов, которые нужно найти
        :return: Список путей к файлам с нужными расширениями
        """
        file_paths = []
        for root, _, files in os.walk(directory):
            for filename in files:
                if filename.lower().endswith(extensions):
                    file_paths.append(os.path.join(root, filename))
        return file_paths

    def run_inventory(self):
        """Метод для создания описи документов."""
        directory = filedialog.askdirectory()  # Выбор каталога
        if not directory:
            logging.warning("Директория не выбрана.")
            messagebox.showwarning("Предупреждение", "Директория не выбрана.")
            return

        self.reference_path.set(directory)  # Сохраняем путь к выбранной директории

        # Инициализируем список для хранения извлечённых данных
        extracted_data = []

        # Получаем все файлы с нужными расширениями во всех подкаталогах
        all_files = self.get_all_files_in_directory(directory)

        # Проходим по каждому файлу и извлекаем данные
        for file_path in all_files:
            ext = os.path.splitext(file_path)[1].lower()
            name, pages = self.extract_metadata(file_path, ext)
            name1 = os.path.splitext(name)[0]  # Убираем расширение из имени файла
            extracted_data.append({
                'name': name,
                'designation': name1,
                'pages': pages,
                'format': ext[1:]
            })

        # Проверяем наличие извлечённых данных перед созданием описи
        if extracted_data:
            # Создаем опись документов
            output_path = os.path.join(directory, "опись.docx")
            try:
                self.create_inventory(extracted_data, output_path)  # Создаем опись
                logging.info("Опись документов успешно создана.")
                messagebox.showinfo("Успех", "Опись документов успешно создана.")
            except Exception as e:
                logging.error(f"Ошибка при создании описи документов: {e}")
                messagebox.showerror("Ошибка", "Ошибка при создании описи документов.")
        else:
            logging.warning("Нет данных для создания описи документов.")
            messagebox.showwarning("Предупреждение", "Нет данных для создания описи документов.")

    def run_apply_numbers(self):
        """Метод для автоматического нанесения номеров на файлы во всех подкаталогах."""
        directory = self.files_directory.get()

        # Пронумеровываем и переименовываем файлы рекурсивно
        index = 1
        for root, _, files in os.walk(directory):
            for file_name in files:
                file_path = os.path.join(root, file_name)
                output_path = os.path.join(root, file_name)  # Сохраняем в той же папке
                # Нанесение текущего номера на файл и сохранение под новым именем
                apply_number_to_file(file_path, index, output_path)
                index += 1

        messagebox.showinfo("Успех", "Номера успешно нанесены на файлы во всех каталогах и подкаталогах.")


def run_rename_files(self):
        rename_file_with_dialog()


if __name__ == "__main__":
    root = tk.Tk()
    app = DocumentProcessorApp(root)
    root.mainloop()